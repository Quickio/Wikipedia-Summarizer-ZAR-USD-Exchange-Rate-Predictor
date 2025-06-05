rom flask import Flask, render_template, request, send_file
import wikipedia
from transformers import pipeline
from fpdf import FPDF
from docx import Document
import os
import tempfile

app = Flask(__name__)
summarizer = pipeline("summarization", model="facebook/bart-large-cnn")

def fetch_wikipedia_summary(topic):
    try:
        wiki_content = wikipedia.page(topic).content
        return wiki_content
    except wikipedia.exceptions.DisambiguationError as e:
        return f"Trt something else: {e.options[:5]}"
    except wikipedia.exceptions.PageError:
        return "Page Error: Topic not found."
    except Exception as e:
        return f"An error occurred: {e}"

def summarize_text(text):
    chunks = []
    while len(text) > 1000:
        split_idx = text[:1000].rfind('.')
        chunks.append(text[:split_idx+1])
        text = text[split_idx+1:]
    chunks.append(text)

    summary = ""
    for chunk in chunks:
        result = summarizer(chunk, max_length=150, min_length=60, do_sample=False)[0]['summary_text']
        summary += result + " "
    return summary.strip()

def save_as_pdf(topic, summary):
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, txt=topic.upper(), ln=True, align='C')
    pdf.ln(10)
    pdf.set_font("Arial", size=12)
    for line in summary.split('\n'):
        pdf.multi_cell(0, 10, line)
    pdf.output(temp_file.name)
    return temp_file.name

def save_as_docx(topic, summary):
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc = Document()
    doc.add_heading(topic.upper(), 0)
    doc.add_paragraph(summary)
    doc.save(temp_file.name)
    return temp_file.name

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        name = request.form["name"]
        topic = request.form["topic"]
        file_type = request.form["filetype"]

        raw_text = fetch_wikipedia_summary(topic)
        if raw_text.startswith("Disambiguation Error") or raw_text.startswith("Page Error") or raw_text.startswith("An error occurred"):
            return render_template("index.html", error=raw_text)

        summary = summarize_text(raw_text[:5000])

        if file_type == "pdf":
            file_path = save_as_pdf(topic, summary)
        else:
            file_path = save_as_docx(topic, summary)

        return send_file(file_path, as_attachment=True)

    return render_template("index.html")


if __name__ == "__main__":
    print("Starting Flask server...")
    app.run(debug=True)


