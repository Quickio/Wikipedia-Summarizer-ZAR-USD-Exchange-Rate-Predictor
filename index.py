import wikipedia
from transformers import pipeline
from fpdf import FPDF
from docx import Document
import tkinter as tk
from tkinter import messagebox, simpledialog


summarizer = pipeline("summarization", model="facebook/bart-large-cnn")

def fetch_wikipedia_summary(topic):
    try:
        wiki_content = wikipedia.page(topic).content
        return wiki_content
    except wikipedia.exceptions.DisambiguationError as e:
        return f"Disambiguation Error: The topic is too broad. Try one of: {e.options[:5]}"
    except wikipedia.exceptions.PageError:
        return "Page Error: Topic not found."
    except Exception as e:
        return f"An error occurred: {e}"

def summarize_text(text):
    chunks = []
    while len(text) > 1000:
        split_idx = text[:1000].rfind('.')
        chunks.append(text[:split_idx+1])
        text = text[split_idx+1:]
    chunks.append(text)

    summary = ""
    for chunk in chunks:
        result = summarizer(chunk, max_length=150, min_length=60, do_sample=False)[0]['summary_text']
        summary += result + " "
    return summary.strip()

def save_as_pdf(topic, summary, name):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, txt=topic.upper(), ln=True, align='C')
    pdf.ln(10)
    pdf.set_font("Arial", size=12)
    for line in summary.split('\n'):
        pdf.multi_cell(0, 10, line)
    filename = f"{topic}_Summary.pdf"
    pdf.output(filename)
    messagebox.showinfo("PDF Created", f"✅ PDF saved as {filename}")

def save_as_docx(topic, summary, name):
    doc = Document()
    doc.add_heading(topic.upper(), 0)
    doc.add_paragraph(summary)
    filename = f"{topic}_Summary.docx"
    doc.save(filename)
    messagebox.showinfo("DOCX Created", f"✅ DOCX saved as {filename}")

def run_summarizer():
    name = simpledialog.askstring("Name", "Enter your name:")
    if not name:
        return

    topic = simpledialog.askstring("Topic", "Enter a Wikipedia topic:")
    if not topic:
        return

    raw_text = fetch_wikipedia_summary(topic)
    if raw_text.startswith("Disambiguation") or raw_text.startswith("Page Error") or raw_text.startswith("An error occurred"):
        messagebox.showerror("Error", raw_text)
        return

    summary = summarize_text(raw_text[:5000]) #limit

    output_choice = simpledialog.askstring("Output Format", "Type 'pdf' or 'docx' to choose output format:")
    if output_choice:
        if output_choice.lower() == "pdf":
            save_as_pdf(topic, summary, name)
        elif output_choice.lower() == "docx":
            save_as_docx(topic, summary, name)
        else:
            messagebox.showwarning("Invalid Input", "Please enter either 'pdf' or 'docx'.")

#gui
root = tk.Tk()
root.title("Wikipedia Topic Summarizer")
root.geometry("400x200")

label = tk.Label(root, text="Wikipedia Summarizer", font=("Arial", 16))
label.pack(pady=20)

run_button = tk.Button(root, text="Start Summarizer", command=run_summarizer, font=("Arial", 12))
run_button.pack(pady=10)

root.mainloop()
